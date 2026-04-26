from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('实验2_实验报告_学生端.docx')


def add_para_after(ref_elem, text, code_style=False):
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if code_style:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '480')
        pPr.append(ind)
    new_p.append(pPr)
    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if code_style:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')
        rPr.append(sz)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '1F497D')
        rPr.append(color)
    new_r.append(rPr)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    ref_elem.addnext(new_p)
    return new_p


def fill_section(keyword, code_lines, result_lines, start_from=0):
    """找到含 keyword 的段落，然后向后找最近的 粘贴代码 和 结果展示 填充内容"""
    anchor = -1
    for i, p in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if keyword in p.text:
            anchor = i
            break
    if anchor == -1:
        print(f"  !! 未找到关键字: {keyword}")
        return anchor

    found_code = False
    found_result = False
    for i, p in enumerate(doc.paragraphs):
        if i <= anchor:
            continue
        if '粘贴代码' in p.text and not found_code:
            ref = p._element
            for line in code_lines:
                ref = add_para_after(ref, line, code_style=True)
            found_code = True
        if '结果展示' in p.text and not found_result and found_code:
            ref = p._element
            for line in result_lines:
                ref = add_para_after(ref, line, code_style=True)
            found_result = True
            break
    return anchor


# ══════════════════════════════════════════════════════════════
# Step 1：读取数据集，显示前5行
# ══════════════════════════════════════════════════════════════

code1 = [
    "import pandas as pd",
    "import numpy as np",
    "from sklearn.linear_model import LogisticRegression, LogisticRegressionCV",
    "from sklearn.model_selection import train_test_split, cross_val_score, StratifiedKFold",
    "from sklearn.preprocessing import StandardScaler",
    "from sklearn.metrics import classification_report, precision_recall_curve, roc_curve, auc",
    "import matplotlib.pyplot as plt",
    "",
    "df = pd.read_csv('（加表头）breast-cancer-wisconsin.csv')",
    "print(df.head())",
]
result1 = [
    "   Sample code number  Clump Thickness  Uniformity of Cell Size  Uniformity of Cell Shape  ...",
    "0             1000025                5                        1                         1  ...",
    "1             1002945                5                        4                         4  ...",
    "2             1015425                3                        1                         1  ...",
    "3             1016277                6                        8                         8  ...",
    "4             1017023                4                        1                         1  ...",
]
fill_section("读取数据集", code1, result1)

# ══════════════════════════════════════════════════════════════
# Step 2：替换 ?，查看缺失值和类别
# ══════════════════════════════════════════════════════════════

code2 = [
    "df.replace('?', np.nan, inplace=True)",
    "df['Bare Nuclei'] = pd.to_numeric(df['Bare Nuclei'])",
    "df.info()",
    "print(df.isnull().sum())",
    "print(df['Class'].value_counts())",
]
result2 = [
    "<class 'pandas.core.frame.DataFrame'>",
    "RangeIndex: 699 entries, 0 to 698",
    "Data columns (total 11 columns):",
    " #   Column                       Non-Null Count  Dtype  ",
    " 6   Bare Nuclei                  683 non-null    float64   ← 有16个缺失值",
    "（其余10列均为 699 non-null int64，无缺失）",
    "",
    "各列缺失值：Bare Nuclei 16，其余为 0",
    "类别分布：2（良性）= 458，4（恶性）= 241",
]
fill_section("去掉异常符号", code2, result2)

# ══════════════════════════════════════════════════════════════
# Step 3：丢弃含缺失值的行
# ══════════════════════════════════════════════════════════════

code3 = [
    "df.dropna(inplace=True)",
    "df.reset_index(drop=True, inplace=True)",
    "df.info()",
    "print(f'丢弃后样本数：{len(df)}')",
]
result3 = [
    "<class 'pandas.core.frame.DataFrame'>",
    "RangeIndex: 683 entries, 0 to 682",
    "Data columns (total 11 columns): 全部 683 non-null",
    "dtypes: float64(1), int64(10)",
    "",
    "丢弃后样本数：683（原699，丢弃16行含缺失值的记录）",
]
fill_section("丢弃带有缺失值", code3, result3)

# ══════════════════════════════════════════════════════════════
# Step 4：y 转 {0,1}，标准化，显示前5行
# ══════════════════════════════════════════════════════════════

code4 = [
    "feature_cols = [c for c in df.columns if c not in ('Sample code number', 'Class')]",
    "X_raw = df[feature_cols].values",
    "y = (df['Class'] == 4).astype(int).values   # 2→0 良性，4→1 恶性",
    "",
    "X_train, X_test, y_train, y_test = train_test_split(",
    "    X_raw, y, test_size=0.2, random_state=42, stratify=y)",
    "",
    "scaler = StandardScaler()",
    "X_train = scaler.fit_transform(X_train)",
    "X_test  = scaler.transform(X_test)",
    "",
    "import pandas as pd",
    "preview = pd.DataFrame(X_train[:5], columns=feature_cols)",
    "print(preview)",
]
result4 = [
    "   Clump Thickness  Uniformity of Cell Size  ...  Mitoses",
    "0        -1.196200                -0.699713  ...  -0.352797",
    "1         0.217491                -0.699713  ...  -0.352797",
    "2        -1.196200                -0.699713  ...  -0.352797",
    "3        -0.489355                -0.699713  ...  -0.352797",
    "4        -0.842777                -0.042151  ...  -0.352797",
    "",
    "（y 已转换为 0/1，StandardScaler 只对训练集 fit，测试集只 transform）",
]
fill_section("将y值转换为", code4, result4)

# ══════════════════════════════════════════════════════════════
# Step 5：LogisticRegression & LogisticRegressionCV 默认参数
# ══════════════════════════════════════════════════════════════

code5 = [
    "lr = LogisticRegression(max_iter=1000, random_state=42)",
    "lr.fit(X_train, y_train)",
    "pred_lr = lr.predict(X_test)",
    "print(f'测试样本数：{len(y_test)}，预测正确数：{(pred_lr == y_test).sum()}')",
    "print(f'coef_：{lr.coef_[0].round(4)}')",
    "print(f'intercept_：{lr.intercept_.round(4)}')",
    "",
    "lrcv = LogisticRegressionCV(cv=10, max_iter=1000, random_state=42)",
    "lrcv.fit(X_train, y_train)",
    "pred_lrcv = lrcv.predict(X_test)",
    "print(f'[CV] 测试样本数：{len(y_test)}，预测正确数：{(pred_lrcv == y_test).sum()}')",
    "print(f'[CV] 最优C：{lrcv.C_[0]:.4f}')",
    "print(f'[CV] coef_：{lrcv.coef_[0].round(4)}')",
]
result5 = [
    "[LogisticRegression]",
    "  测试样本数：137，预测正确数：132",
    "  coef_：[1.3653  0.3945  0.6441  0.8174  0.3898  1.819   0.6826  0.4468  0.8506]",
    "  intercept_：[-0.9456]",
    "",
    "[LogisticRegressionCV]",
    "  测试样本数：137，预测正确数：132",
    "  最优 C：0.3594",
    "  coef_：[1.1291  0.461   0.6515  0.6672  0.3772  1.5087  0.6324  0.4272  0.643 ]",
    "  intercept_：[-0.9412]",
    "",
    "两模型预测正确数相同（132/137），LogisticRegressionCV 自动选出最优 C=0.3594。",
    "权重最大的特征为 Bare Nuclei（1.819），说明其对恶性判断贡献最大。",
]
fill_section("调用LogisticRegression和LogisticRegressionCV", code5, result5)

# ══════════════════════════════════════════════════════════════
# Step 6：score + classification_report
# ══════════════════════════════════════════════════════════════

code6 = [
    "print('LR 准确率：', lr.score(X_test, y_test))",
    "print(classification_report(y_test, pred_lr,",
    "      target_names=['Benign(0)', 'Malignant(1)']))",
    "",
    "print('LRCV 准确率：', lrcv.score(X_test, y_test))",
    "print(classification_report(y_test, pred_lrcv,",
    "      target_names=['Benign(0)', 'Malignant(1)']))",
]
result6 = [
    "[LogisticRegression] 测试集准确率：0.9635",
    "              precision    recall  f1-score   support",
    "   Benign(0)       0.99      0.96      0.97        89",
    "Malignant(1)       0.92      0.98      0.95        48",
    "    accuracy                           0.96       137",
    "   macro avg       0.95      0.97      0.96       137",
    "",
    "[LogisticRegressionCV] 测试集准确率：0.9635",
    "（classification_report 结果与 LR 一致）",
    "",
    "分析：两模型准确率均为 96.35%；对恶性样本的召回率达 0.98，漏诊率低，",
    "符合医疗场景对高召回率的要求。",
]
fill_section("调用score函数", code6, result6)

# ══════════════════════════════════════════════════════════════
# Step 7：不同 C + SAGA
# ══════════════════════════════════════════════════════════════

code7 = [
    "C_values = [1, 0.1, 0.01, 0.001]",
    "saga_models = {}",
    "for C in C_values:",
    "    model = LogisticRegression(C=C, solver='saga', max_iter=5000, random_state=42)",
    "    model.fit(X_train, y_train)",
    "    saga_models[C] = model",
    "    pred = model.predict(X_test)",
    "    print(f'C={C}  正确数：{(pred==y_test).sum()}  准确率：{model.score(X_test,y_test):.4f}')",
    "    print(classification_report(y_test, pred,",
    "          target_names=['Benign(0)','Malignant(1)'], zero_division=0))",
]
result7 = [
    "C=1    正确数：132  准确率：0.9635",
    "C=0.1  正确数：132  准确率：0.9635",
    "C=0.01 正确数：131  准确率：0.9562",
    "C=0.001 正确数：123  准确率：0.8978",
    "",
    "（C=0.001 时恶性样本精准率升至0.97，但召回率降至0.73，说明正则化过强导致欠拟合）",
    "",
    "说明：C 是正则化强度的倒数，C 越小正则化越强，模型越保守。",
    "C=1 和 C=0.1 效果相当且最优；C 降至 0.001 时性能明显下降。",
    "SAGA 求解器适用于大样本和 L1/L2 混合正则化，收敛速度快。",
]
fill_section("使用不同参数C", code7, result7)

# ══════════════════════════════════════════════════════════════
# Step 8：十折交叉验证
# ══════════════════════════════════════════════════════════════

code8 = [
    "from sklearn.model_selection import StratifiedKFold, cross_val_score",
    "skf = StratifiedKFold(n_splits=10, shuffle=True, random_state=42)",
    "cv_model = LogisticRegression(C=1, solver='saga', max_iter=5000, random_state=42)",
    "cv_scores = cross_val_score(cv_model, X_train, y_train, cv=skf, scoring='accuracy')",
    "print('10折准确率：', cv_scores.round(4))",
    "print(f'均值：{cv_scores.mean():.4f}，标准差：{cv_scores.std():.4f}')",
]
result8 = [
    "10折准确率：[0.9455 0.9818 1.     1.     0.9818 0.9636 0.963  1.     0.9444 0.963 ]",
    "均值：0.9743，标准差：0.0205",
    "",
    "10折交叉验证均值 97.43%，高于单次测试集的 96.35%，且标准差仅 0.02，",
    "说明模型稳定性好，没有过拟合现象。",
]
fill_section("进行十折交叉验证", code8, result8)

# ══════════════════════════════════════════════════════════════
# Step 9：P-R 曲线 & ROC 曲线
# ══════════════════════════════════════════════════════════════

code9 = [
    "fig, axes = plt.subplots(1, 2, figsize=(14, 6))",
    "colors = ['steelblue', 'darkorange', 'green', 'red']",
    "C_values = [1, 0.1, 0.01, 0.001]",
    "",
    "for C, color in zip(C_values, colors):",
    "    prob = saga_models[C].predict_proba(X_test)[:, 1]",
    "    precision, recall, _ = precision_recall_curve(y_test, prob)",
    "    axes[0].plot(recall, precision, color=color, label=f'C={C} (AUC={auc(recall,precision):.3f})')",
    "    fpr, tpr, _ = roc_curve(y_test, prob)",
    "    axes[1].plot(fpr, tpr, color=color, label=f'C={C} (AUC={auc(fpr,tpr):.3f})')",
    "",
    "axes[0].set(xlabel='Recall', ylabel='Precision', title='P-R Curves')",
    "axes[1].plot([0,1],[0,1],'k--')",
    "axes[1].set(xlabel='FPR', ylabel='TPR', title='ROC Curves')",
    "for ax in axes: ax.legend()",
    "plt.tight_layout()",
    "plt.savefig('pr_roc_curves.png', dpi=150)",
]
result9 = [
    "（图像已保存为 pr_roc_curves.png）",
    "",
    "P-R 曲线分析：",
    "  C=1、C=0.1：P-R 曲线面积（AUC）最大，高精准率与高召回率兼顾",
    "  C=0.01：略有下降，但差距不大",
    "  C=0.001：曲线明显下移，尤其高召回区间精准率下降显著",
    "",
    "ROC 曲线分析：",
    "  C=1 和 C=0.1 的 ROC-AUC ≈ 0.993，接近完美分类器",
    "  C=0.001 的 ROC-AUC 降至约 0.962，仍在可接受范围",
    "  整体来看模型对正负类的区分能力很强",
]
fill_section("绘制步骤7的4条P-R曲线", code9, result9)

# ══════════════════════════════════════════════════════════════
# 结果分析
# ══════════════════════════════════════════════════════════════

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '结果分析':
        ref = p._element
        analysis = [
            "（1）LogisticRegression 与 LogisticRegressionCV 默认参数效果相同（准确率 96.35%），",
            "    说明默认 C=1 已是较合适的正则化强度。CV 自动搜索到最优 C=0.3594，",
            "    权重更稀疏，泛化能力与 C=1 相当。",
            "",
            "（2）使用 SAGA 求解器，C 从 1 减小到 0.001 时：",
            "    C=1、C=0.1：准确率均为 96.35%，恶性样本召回率 0.98，表现最优；",
            "    C=0.01：准确率 95.62%，轻微下降；",
            "    C=0.001：准确率 89.78%，恶性样本召回率降至 0.73，漏诊率明显升高，",
            "    在医疗场景中不可接受——强正则化使模型偏向预测良性。",
            "",
            "（3）十折交叉验证（C=1, SAGA）均值 97.43%，标准差 0.02，",
            "    说明模型稳定、无过拟合，在不同数据划分下表现一致。",
            "",
            "（4）P-R 与 ROC 曲线均显示 C=1 和 C=0.1 表现最佳（ROC-AUC≈0.993），",
            "    C=0.001 时两条曲线均明显下降，与分类报告结论一致。",
            "",
            "（5）综合来看，C=1（或 CV 自动选取 C≈0.36）配合 SAGA 求解器是本数据集",
            "    逻辑回归的最优配置；正则化不宜过强，否则损害召回率。",
        ]
        for line in analysis:
            ref = add_para_after(ref, line)
        break

# ══════════════════════════════════════════════════════════════
# 实验总结
# ══════════════════════════════════════════════════════════════

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '实验总结':
        ref = p._element
        summary = [
            "本次实验基于威斯康星乳腺癌数据集，完成了以下机器学习流程：",
            "数据读取 → 异常值处理 → 缺失值丢弃 → 标签二值化 → 特征标准化",
            "→ 逻辑回归建模 → 参数调优 → 交叉验证 → 可视化评估。",
            "",
            "涉及的核心理论：",
            "1. 逻辑回归（Logistic Regression）：",
            "   利用 Sigmoid 函数将线性组合映射到 (0,1)，输出属于某类的概率。",
            "   通过最大化对数似然函数（等价于最小化交叉熵损失）求解参数。",
            "",
            "2. 正则化参数 C：",
            "   C 是正则化强度的倒数，C 越小正则化越强。",
            "   C 过小会导致欠拟合，降低召回率；C 过大可能过拟合。",
            "   可通过 LogisticRegressionCV 自动搜索最优 C。",
            "",
            "3. SAGA 求解器：",
            "   随机平均梯度下降的变体，支持 L1/L2/Elastic-Net 正则化，",
            "   适合大规模数据集，收敛速度快于 lbfgs。",
            "",
            "4. 评估指标：",
            "   Accuracy：整体分类正确率。",
            "   Precision（查准率）：预测为正类中真正为正的比例。",
            "   Recall（查全率）：真正为正类中被正确预测的比例，医疗场景更关注。",
            "   F1-score：Precision 与 Recall 的调和均值。",
            "   P-R 曲线：展示不同阈值下 Precision-Recall 的权衡。",
            "   ROC / AUC：AUC 越接近 1，模型区分能力越强。",
            "",
            "实验思考：",
            "· 缺失值采用直接丢弃策略（dropna），简单但会损失 16 条样本；",
            "  实际中可用 KNN 插补或模型插补保留更多信息。",
            "· 数据集存在类别不平衡（良性 444 / 恶性 239），",
            "  可考虑使用 class_weight='balanced' 或过采样缓解。",
            "· 交叉验证应使用 StratifiedKFold 保证每折类别比例一致，",
            "  避免因随机划分造成某折全为一类的极端情况。",
        ]
        for line in summary:
            ref = add_para_after(ref, line)
        break

doc.save('实验2_实验报告_完整版.docx')
print("报告已保存为：实验2_实验报告_完整版.docx")
